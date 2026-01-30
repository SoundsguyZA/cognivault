#!/usr/bin/env python3
"""
Document Processor for CogniVault
Handles text extraction from various document formats
"""

import json
from pathlib import Path
from typing import Optional, Dict, Any
import tempfile
import subprocess
import shutil

class DocumentProcessor:
    def __init__(self):
        """Initialize document processor"""
        self.supported_formats = {'.txt', '.md', '.json', '.pdf', '.docx', '.html', '.rtf'}
        
        # Check available tools
        self.tools_available = self.check_available_tools()
    
    def check_available_tools(self) -> Dict[str, bool]:
        """Check which document processing tools are available"""
        tools = {}
        
        # Check for pdfplumber (PDF processing)
        try:
            import pdfplumber
            tools['pdfplumber'] = True
        except ImportError:
            tools['pdfplumber'] = False
        
        # Check for python-docx (DOCX processing)
        try:
            import docx
            tools['python_docx'] = True
        except ImportError:
            tools['python_docx'] = False
        
        # Check for BeautifulSoup (HTML processing)
        try:
            from bs4 import BeautifulSoup
            tools['beautifulsoup'] = True
        except ImportError:
            tools['beautifulsoup'] = False
        
        # Check for pandoc (universal converter)
        tools['pandoc'] = shutil.which('pandoc') is not None
        
        return tools
    
    def process_document(self, file_path: Path) -> Optional[str]:
        """Process document and extract text content"""
        if not file_path.exists():
            print(f"File not found: {file_path}")
            return None
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            print(f"Unsupported format: {file_extension}")
            return None
        
        try:
            if file_extension == '.txt':
                return self.process_text_file(file_path)
            elif file_extension == '.md':
                return self.process_markdown_file(file_path)
            elif file_extension == '.json':
                return self.process_json_file(file_path)
            elif file_extension == '.pdf':
                return self.process_pdf_file(file_path)
            elif file_extension == '.docx':
                return self.process_docx_file(file_path)
            elif file_extension == '.html':
                return self.process_html_file(file_path)
            elif file_extension == '.rtf':
                return self.process_rtf_file(file_path)
            else:
                return self.fallback_text_extraction(file_path)
        
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return self.fallback_text_extraction(file_path)
    
    def process_text_file(self, file_path: Path) -> str:
        """Process plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    print(f"Successfully read text file with {encoding} encoding")
                    return content
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, 'rb') as f:
                content = f.read().decode('utf-8', errors='ignore')
            
            return content
        except Exception as e:
            print(f"Error reading text file: {e}")
            return f"Error reading file: {file_path.name}"
    
    def process_markdown_file(self, file_path: Path) -> str:
        """Process Markdown file"""
        content = self.process_text_file(file_path)
        
        # Add some basic markdown metadata
        metadata = f"Markdown file: {file_path.name}\n"
        metadata += f"Length: {len(content)} characters\n"
        metadata += "=" * 50 + "\n\n"
        
        return metadata + content
    
    def process_json_file(self, file_path: Path) -> str:
        """Process JSON file and make it searchable"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to searchable text
            searchable_text = self.json_to_searchable_text(data, file_path.name)
            return searchable_text
        
        except json.JSONDecodeError as e:
            print(f"Invalid JSON in {file_path}: {e}")
            # Fall back to treating as text
            return self.process_text_file(file_path)
        except Exception as e:
            print(f"Error processing JSON file: {e}")
            return f"Error processing JSON file: {file_path.name}"
    
    def json_to_searchable_text(self, data: Any, filename: str, prefix: str = "") -> str:
        """Convert JSON data to searchable text format"""
        lines = [f"JSON file: {filename}", "=" * 50, ""]
        
        def extract_text(obj, current_prefix=""):
            if isinstance(obj, dict):
                for key, value in obj.items():
                    new_prefix = f"{current_prefix}.{key}" if current_prefix else key
                    if isinstance(value, (dict, list)):
                        lines.append(f"{new_prefix}:")
                        extract_text(value, new_prefix)
                    else:
                        lines.append(f"{new_prefix}: {value}")
            
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    new_prefix = f"{current_prefix}[{i}]" if current_prefix else f"item_{i}"
                    if isinstance(item, (dict, list)):
                        lines.append(f"{new_prefix}:")
                        extract_text(item, new_prefix)
                    else:
                        lines.append(f"{new_prefix}: {item}")
            else:
                lines.append(f"{current_prefix}: {obj}")
        
        extract_text(data)
        return "\n".join(lines)
    
    def process_pdf_file(self, file_path: Path) -> str:
        """Process PDF file"""
        if self.tools_available['pdfplumber']:
            return self.process_pdf_with_pdfplumber(file_path)
        elif self.tools_available['pandoc']:
            return self.process_with_pandoc(file_path)
        else:
            return self.install_and_process_pdf(file_path)
    
    def process_pdf_with_pdfplumber(self, file_path: Path) -> str:
        """Process PDF using pdfplumber"""
        try:
            import pdfplumber
            
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                text_content.append(f"PDF file: {file_path.name}")
                text_content.append(f"Pages: {len(pdf.pages)}")
                text_content.append("=" * 50)
                text_content.append("")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num} ---")
                        text_content.append(page_text)
                        text_content.append("")
            
            return "\n".join(text_content)
        
        except Exception as e:
            print(f"Error processing PDF with pdfplumber: {e}")
            return self.fallback_text_extraction(file_path)
    
    def install_and_process_pdf(self, file_path: Path) -> str:
        """Install pdfplumber and process PDF"""
        try:
            import subprocess
            import sys
            
            print("Installing pdfplumber for PDF processing...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", "pdfplumber"])
            
            # Update tools availability
            self.tools_available['pdfplumber'] = True
            
            # Try processing again
            return self.process_pdf_with_pdfplumber(file_path)
        
        except Exception as e:
            print(f"Failed to install pdfplumber: {e}")
            return self.fallback_text_extraction(file_path)
    
    def process_docx_file(self, file_path: Path) -> str:
        """Process DOCX file"""
        if self.tools_available['python_docx']:
            return self.process_docx_with_python_docx(file_path)
        else:
            return self.install_and_process_docx(file_path)
    
    def process_docx_with_python_docx(self, file_path: Path) -> str:
        """Process DOCX using python-docx"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            text_content = []
            text_content.append(f"DOCX file: {file_path.name}")
            text_content.append(f"Paragraphs: {len(doc.paragraphs)}")
            text_content.append("=" * 50)
            text_content.append("")
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract table content if any
            if doc.tables:
                text_content.append("\n--- Tables ---")
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            text_content.append(row_text)
            
            return "\n".join(text_content)
        
        except Exception as e:
            print(f"Error processing DOCX: {e}")
            return self.fallback_text_extraction(file_path)
    
    def install_and_process_docx(self, file_path: Path) -> str:
        """Install python-docx and process DOCX"""
        try:
            import subprocess
            import sys
            
            print("Installing python-docx for DOCX processing...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            
            # Update tools availability
            self.tools_available['python_docx'] = True
            
            # Try processing again
            return self.process_docx_with_python_docx(file_path)
        
        except Exception as e:
            print(f"Failed to install python-docx: {e}")
            return self.fallback_text_extraction(file_path)
    
    def process_html_file(self, file_path: Path) -> str:
        """Process HTML file"""
        if self.tools_available['beautifulsoup']:
            return self.process_html_with_bs4(file_path)
        else:
            return self.process_html_basic(file_path)
    
    def process_html_with_bs4(self, file_path: Path) -> str:
        """Process HTML using BeautifulSoup"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            metadata = f"HTML file: {file_path.name}\n"
            metadata += f"Title: {soup.title.string if soup.title else 'No title'}\n"
            metadata += "=" * 50 + "\n\n"
            
            return metadata + text
        
        except Exception as e:
            print(f"Error processing HTML with BeautifulSoup: {e}")
            return self.process_html_basic(file_path)
    
    def process_html_basic(self, file_path: Path) -> str:
        """Basic HTML processing without BeautifulSoup"""
        import re
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Remove HTML tags using regex (basic approach)
            clean_text = re.sub(r'<[^>]+>', '', html_content)
            
            # Clean up extra whitespace
            clean_text = re.sub(r'\s+', ' ', clean_text).strip()
            
            metadata = f"HTML file: {file_path.name}\n"
            metadata += "=" * 50 + "\n\n"
            
            return metadata + clean_text
        
        except Exception as e:
            print(f"Error processing HTML: {e}")
            return self.fallback_text_extraction(file_path)
    
    def process_rtf_file(self, file_path: Path) -> str:
        """Process RTF file"""
        if self.tools_available['pandoc']:
            return self.process_with_pandoc(file_path)
        else:
            return self.fallback_text_extraction(file_path)
    
    def process_with_pandoc(self, file_path: Path) -> str:
        """Process document using pandoc"""
        try:
            cmd = ['pandoc', str(file_path), '-t', 'plain']
            result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')
            
            if result.returncode == 0:
                metadata = f"Document: {file_path.name} (processed with pandoc)\n"
                metadata += "=" * 50 + "\n\n"
                return metadata + result.stdout
            else:
                print(f"Pandoc error: {result.stderr}")
                return self.fallback_text_extraction(file_path)
        
        except Exception as e:
            print(f"Error using pandoc: {e}")
            return self.fallback_text_extraction(file_path)
    
    def fallback_text_extraction(self, file_path: Path) -> str:
        """Fallback text extraction for unsupported formats"""
        try:
            # Try to read as text with error handling
            with open(file_path, 'rb') as f:
                raw_content = f.read()
            
            # Decode with error handling
            text_content = raw_content.decode('utf-8', errors='ignore')
            
            # Clean up non-printable characters
            import string
            printable = set(string.printable)
            clean_content = ''.join(filter(lambda x: x in printable, text_content))
            
            metadata = f"File: {file_path.name} (fallback extraction)\n"
            metadata += f"Original size: {len(raw_content)} bytes\n"
            metadata += f"Extracted text: {len(clean_content)} characters\n"
            metadata += "=" * 50 + "\n\n"
            
            return metadata + clean_content[:10000]  # Limit to first 10K chars
        
        except Exception as e:
            return f"Unable to process file: {file_path.name}\nError: {str(e)}"
    
    def get_supported_formats(self) -> set:
        """Get list of supported document formats"""
        return self.supported_formats.copy()
    
    def is_supported_format(self, file_path: Path) -> bool:
        """Check if file format is supported"""
        return file_path.suffix.lower() in self.supported_formats
    
    def get_document_info(self, file_path: Path) -> Dict[str, Any]:
        """Get document information"""
        if not file_path.exists():
            return {}
        
        stat = file_path.stat()
        
        return {
            'filename': file_path.name,
            'path': str(file_path),
            'size': stat.st_size,
            'modified': stat.st_mtime,
            'extension': file_path.suffix.lower(),
            'supported': self.is_supported_format(file_path),
            'available_tools': self.tools_available
        }